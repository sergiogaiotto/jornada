"""Gerador do documento executivo do portão GO via python-docx (§8-M4-A3, §3.2).

Implementa GeradorDocumentoPort (§2.1). Determinístico e em memória (BytesIO):
nenhum LLM, nenhum filesystem — os bytes vão direto para `documento_portao` (§4.1).
"""

from io import BytesIO
from typing import Any

from docx import Document

from domain.campanha.modelos import OS, Pendencia
from domain.validacao.modelos import ValidacaoCampo


class GeradorDocxPortao:
    def documento_go(
        self,
        *,
        os_: OS,
        validacoes: list[ValidacaoCampo],
        pendencias: list[Pendencia],
        frozen: dict[str, Any],
    ) -> bytes:
        doc = Document()
        doc.add_heading(f"Documento Executivo — Portão GO · {os_.codigo}", level=1)
        doc.add_paragraph(f"Campanha: {os_.nome} (t-shirt {os_.tshirt})")
        doc.add_paragraph(f"Fase após o GO: criada · Congelado em: {frozen.get('congelado_em')}")

        doc.add_heading("Briefing validado (campo a campo)", level=2)
        ultima_por_campo = {v.campo: v for v in validacoes}
        for campo, entrada in os_.briefing.items():
            valor = entrada.get("valor") if isinstance(entrada, dict) else entrada
            validacao = ultima_por_campo.get(campo)
            veredito = validacao.veredito if validacao else "decidido via pendência"
            fonte = validacao.evidencia.get("fonte") if validacao else None
            sufixo = f" · fonte: {fonte}" if fonte else ""
            doc.add_paragraph(f"{campo}: {valor} — {veredito}{sufixo}", style="List Bullet")

        doc.add_heading("Pendências", level=2)
        if pendencias:
            for pendencia in pendencias:
                doc.add_paragraph(
                    f"#{pendencia.numero} {pendencia.titulo} — {pendencia.status}"
                    + (f" ({pendencia.severidade})" if pendencia.severidade else ""),
                    style="List Bullet",
                )
        else:
            doc.add_paragraph("Nenhuma pendência registrada.")

        doc.add_heading("Versões congeladas (os.frozen §4.1)", level=2)
        for agente, versao in sorted(dict(frozen.get("agent_versions", {})).items()):
            doc.add_paragraph(f"Agente {agente}: v{versao}", style="List Bullet")
        doc.add_paragraph(f"Política: v{frozen.get('policy_version')}", style="List Bullet")
        doc.add_paragraph(f"Tarifário: {frozen.get('tarifario_id')}", style="List Bullet")

        doc.add_heading("SLAs congelados", level=2)
        for etapa, sla in dict(frozen.get("slas", {})).items():
            doc.add_paragraph(
                f"{etapa}: {sla.get('sla_dias')} dias — prazo {sla.get('prazo')}",
                style="List Bullet",
            )

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def relatorio_segmento(self, *, relatorio: dict[str, Any]) -> bytes:
        """.docx do relatório T5a (§8-M5): funil bruto→elegível→líquido, waterfall,
        sobreposição, volume de abordagem por canal e frescor por fonte. Determinístico."""
        segmento = dict(relatorio.get("segmento", {}))
        funil = dict(relatorio.get("funil", {}))

        doc = Document()
        doc.add_heading(f"Relatório de Segmento Data Cloud · {segmento.get('id')}", level=1)
        doc.add_paragraph(f"{segmento.get('nome')} — {segmento.get('criterios_resumo')}")
        doc.add_paragraph(
            f"Ciclo: {segmento.get('ciclo')} · Status: {segmento.get('status')} · "
            f"Republicado em: {segmento.get('republicado_em')}"
        )

        doc.add_heading("Funil (bruto → elegível → líquido)", level=2)
        for etapa in ("bruto", "elegivel", "liquido"):
            doc.add_paragraph(f"{etapa}: {funil.get(etapa)}", style="List Bullet")

        doc.add_heading("Waterfall", level=2)
        for corte in list(relatorio.get("waterfall", [])):
            doc.add_paragraph(
                f"{corte.get('etapa')}: -{corte.get('corte')} → {corte.get('restante')} "
                f"({corte.get('motivo')})",
                style="List Bullet",
            )

        doc.add_heading("Sobreposição com outros segmentos", level=2)
        sobreposicao = dict(relatorio.get("sobreposicao", {}))
        if sobreposicao:
            for outro, contatos in sobreposicao.items():
                doc.add_paragraph(f"{outro}: {contatos} contatos", style="List Bullet")
        else:
            doc.add_paragraph("Sem sobreposição registrada.")

        doc.add_heading("Volume de abordagem por canal (pós caps/quiet/colisões)", level=2)
        colisoes = dict(relatorio.get("colisoes", {}))
        for canal, volume in dict(relatorio.get("volume_abordagem", {})).items():
            doc.add_paragraph(
                f"{canal}: {volume.get('n')} ({volume.get('pct')}% do líquido) · "
                f"colisões (governor): {colisoes.get(canal, 0)}",
                style="List Bullet",
            )

        doc.add_heading("Frescor por fonte", level=2)
        for fonte, ultima in dict(relatorio.get("frescor", {})).items():
            doc.add_paragraph(f"{fonte}: {ultima}", style="List Bullet")

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
